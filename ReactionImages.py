import json
import os
from rdkit import Chem
from rdkit.Chem import Draw, rdChemReactions
from PIL import Image
from docx import Document
from docx.shared import Inches

# # Create a molecule from SMILES
# mol = Chem.MolFromSmiles('c1ccncc1')

# # Create a reaction from SMARTS
# rxn = Reactions.ReactionFromSmarts('Nc1ncn[nH]1.OB(O)c1ccccc1CBr>>BrCc1ccccc1Nc1ncn[nH]1', useSmiles=True)

# # Generate the molecule image
# mol_img = Draw.MolToImage(mol)

# # Generate the reaction image
# rxn_img = Draw.ReactionToImage(rxn)

# # Save the molecule image
# mol_img.save("molecule.png", format="PNG")

# # Save the reaction image
# rxn_img.save("reaction.png", format="PNG")

# print("Images saved as molecule.png and reaction.png")

# def generate_reaction_image(reaction_smarts, image_path):
#     """Generate and save an image of the reaction."""
#     rxn = rdChemReactions.ReactionFromSmarts(reaction_smarts, useSmiles=True)
#     img = Draw.ReactionToImage(rxn)
#     img.save(image_path, format='PNG')

def generate_reaction_image(reaction_smarts, image_path):
    """
    Saves image of reaction SMILES.
    Returns True if image is successfully generated and False otherwise.
    """
    try:
        rxn = rdChemReactions.ReactionFromSmarts(reaction_smarts, useSmiles=True)
        img = Draw.ReactionToImage(rxn)
        img.save(image_path, format='PNG')
        return True
    except Exception as e:
        print(f"Error generating image for reaction SMARTS {reaction_smarts}: {e}")
        return False

def create_document(data, word_doc):
    """Create a Word document with reaction drawings."""
    # Create a new Word document
    doc = Document()
    doc.add_heading('Synthesis Pathways', level=1)

    count = 1

    # Iterate over the JSON data
    for idx, (final_product, pathway) in enumerate(data.items(), start=1):
        # Add a header for the final product
        doc.add_heading(f'Product {idx}', level=2)

        print(f'on product {count}')
        count += 1

        # Process each step in the pathway
        for num, step in enumerate(pathway):
            reaction_smarts = step
            image_path = os.path.join(image_dir, f'{final_product}_step{num}.png')

            # Generate and save the reaction image
            image_saved = generate_reaction_image(reaction_smarts, image_path)

            if image_saved:
                # Add image to Word document
                # doc.add_paragraph()  # Add a description or label for the image
                doc.add_picture(image_path, width=Inches(6))  # Adjust width as needed

            else:
                doc.add_paragraph(reaction_smarts)

    # Save the Word document
    doc.save(f'{word_doc}.docx')
    print("Word document saved as synthesis_pathways.docx")

if __name__ == "__main__":
    # Load JSON data
    with open('/Users/angelinaning/Downloads/jensen_lab_urop/reaction_pathways/reaction_pathways_code/mols_iter2/iter2_synthesis_pathways.json', 'r') as f:
        data = json.load(f)

    # Directory to save images
    image_dir = 'reaction_images'
    os.makedirs(image_dir, exist_ok=True)

    # mol = Chem.MolFromSmiles('O=[N+]([O-])c1cccnc1Br')
    # mol_img = Draw.MolToImage(mol)
    # mol_img.save("moledcule.png", format='PNG')

    # create_document(data, 'iter2_synthesis_pathways')
    # image_path = 'reaction_images/CCC/C=C/CN(c1ccc(OC)c(OC)c1)c1cc(-c2ccc(Cl)cc2)cnc1N_step0.png'
    # generate_reaction_image('C1=C(C=NC(=C1N)N)Br.OB(O)c1ccc(Cl)cc1>>Nc1cc(-c2ccc(Cl)cc2)cnc1N', image_path)
    # print('image saved')
